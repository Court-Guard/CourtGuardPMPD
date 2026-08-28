import os
from core.exceptions import ExtractionFailedError
from ingestion.extractors.document_extractor import DocumentExtractor

try:
    import docx
except ImportError:
    docx = None

class WordExtractor(DocumentExtractor):
    """
    Extracts structured text from a Microsoft Word (.docx) document
    using the python-docx library.
    """
    def extract_text(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise ExtractionFailedError(f"Word document not found: {file_path}")
        if docx is None:
            raise ExtractionFailedError("python-docx library is not installed. Run: pip install python-docx")
            
        try:
            doc = docx.Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract basic table structure if present
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_data))
                    
            return "\n".join(full_text)
        except Exception as e:
            raise ExtractionFailedError(f"Failed to extract Word document: {str(e)}")
